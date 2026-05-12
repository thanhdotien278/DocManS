from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/Super/DocManS")
OUTPUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUTPUT_DIR / "assets"
TMP_DIR = ROOT / "tmp" / "docs"
DOCX_PATH = OUTPUT_DIR / "bao-cao-trinh-giam-doc-rtms.docx"

PRIMARY = RGBColor(20, 90, 55)
ACCENT = RGBColor(167, 124, 31)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)


def set_default_font(document):
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT


def add_title(document, text, subtitle=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = PRIMARY

    if subtitle:
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(11)
        run2.font.color.rgb = MUTED


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(16 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(document, text, bold_prefix=None):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
      prefix = p.add_run(bold_prefix)
      prefix.bold = True
      prefix.font.name = "Arial"
      prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
      prefix.font.color.rgb = TEXT
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = TEXT
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = TEXT
    return p


def add_caption(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_image(document, path, width_cm, caption):
    document.add_picture(str(path), width=Cm(width_cm))
    add_caption(document, caption)


def build_overview_table(document):
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)

    rows = [
        ("Tên đề tài", "Hệ thống quản lý NCKH, CN và đổi mới sáng tạo (RTMS / DocManSystem)"),
        ("Đơn vị sử dụng", "Học viện Quân y"),
        ("Mục tiêu giai đoạn 1", "Số hóa quy trình tiếp nhận, phê duyệt, theo dõi đề tài, giao việc và dashboard điều hành"),
        ("Hiện trạng kỹ thuật", "Đã có nền tảng workspace, web app, auth nội bộ và bộ giao diện nghiệp vụ đầu tiên"),
    ]

    for row_idx, (label, value) in enumerate(rows):
        left = table.cell(row_idx, 0)
        right = table.cell(row_idx, 1)
        left.text = label
        right.text = value
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, "EAF3EE")
        for paragraph in left.paragraphs + right.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(10.5)
        for run in left.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = PRIMARY


def build_stack_table(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Lớp", "Công nghệ", "Lý do lựa chọn"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        set_cell_shading(cell, "145A37")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    rows = [
        ("Frontend", "Next.js, React, TypeScript, Tailwind CSS", "Phù hợp web nội bộ nhiều màn quản trị, dễ tái sử dụng component, mở rộng tốt cho dashboard và biểu mẫu."),
        ("Backend", "NestJS, TypeScript", "Tách module nghiệp vụ rõ ràng, dễ áp dụng guard, validation, audit log và quy tắc workflow."),
        ("Dữ liệu", "PostgreSQL, Prisma", "Phù hợp dữ liệu quan hệ, bảo đảm tính nhất quán, migration rõ ràng, dễ kiểm soát schema."),
        ("Phiên và nền tảng mở rộng", "Redis", "Dùng cho cache, queue, nhắc việc, notification jobs ở các bước tiếp theo."),
        ("Tệp đính kèm", "MinIO", "Phù hợp lưu hồ sơ, báo cáo và tài liệu đính kèm theo mô hình object storage nội bộ."),
        ("Triển khai", "Docker Compose, Nginx", "Triển khai đơn giản cho giai đoạn 1, dễ vận hành nội bộ và kiểm soát hạ tầng."),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(10)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    set_default_font(document)
    section = document.sections[0]
    set_page_margins(section)

    logo_path = ROOT / "apps" / "web" / "public" / "logo.png"
    if logo_path.exists():
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(logo_path), width=Cm(2.4))

    add_title(
        document,
        "BÁO CÁO TÓM TẮT ĐỀ TÀI TRIỂN KHAI HỆ THỐNG RTMS",
        "Phục vụ trình bày với Giám đốc"
    )

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Ngày lập báo cáo: 03/05/2026")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = MUTED

    add_heading(document, "1. Tóm tắt điều hành")
    add_body(
        document,
        "Đề tài RTMS nhằm xây dựng một hệ thống quản lý tập trung cho toàn bộ vòng đời đề tài nghiên cứu khoa học cấp trường, thay thế cách quản lý phân tán bằng Excel, email và hồ sơ rời rạc. Ở giai đoạn hiện tại, project đã hình thành nền tảng kỹ thuật, xác thực nội bộ và bộ giao diện nghiệp vụ đầu tiên để phục vụ trình diễn và phát triển tiếp theo."
    )
    build_overview_table(document)

    add_heading(document, "2. Lý do nên triển khai đề tài")
    reasons = [
        "Chuẩn hóa một quy trình nghiệp vụ đang phân tán trên nhiều công cụ, giúp giảm thất lạc hồ sơ, bỏ sót việc và chậm phê duyệt.",
        "Tạo khả năng theo dõi tập trung cho lãnh đạo: nhìn thấy hồ sơ chờ duyệt, đề tài chậm tiến độ, việc quá hạn và cảnh báo ưu tiên trên một dashboard thống nhất.",
        "Tăng tính minh bạch và truy vết nhờ cơ chế trạng thái xử lý, timeline nghiệp vụ và audit log cho các thao tác quan trọng.",
        "Tạo nền tảng để mở rộng các chức năng cốt lõi của Học viện như quản lý tiến độ, giao việc, thông báo nhắc việc, báo cáo và xuất dữ liệu.",
        "Phù hợp mô hình vận hành nội bộ của nhà trường: quản lý theo vai trò, theo đơn vị và theo phạm vi dữ liệu, thay vì mô hình SaaS chung chung."
    ]
    for reason in reasons:
        add_bullet(document, reason)

    add_heading(document, "3. Tech stack đề xuất và đã chốt")
    add_body(
        document,
        "Tech stack hiện tại được chốt theo hướng modular monolith để dễ triển khai, dễ vận hành nội bộ, nhưng vẫn đủ khả năng mở rộng khi số lượng quy trình và người dùng tăng lên."
    )
    build_stack_table(document)

    add_heading(document, "4. Những bước đã triển khai")
    completed = [
        "Khởi tạo workspace kỹ thuật theo mô hình monorepo với `apps/web`, `apps/api` và các package dùng chung.",
        "Thiết lập web app quản trị với shell tổng thể gồm sidebar, topbar, breadcrumb, tìm kiếm nhanh và ngữ cảnh người dùng.",
        "Xây dựng các màn hình nghiệp vụ đầu tiên để trình diễn luồng sử dụng: đăng nhập, dashboard lãnh đạo, danh sách hồ sơ đề tài, chi tiết hồ sơ, danh sách giao việc.",
        "Triển khai auth nội bộ phase 1: đăng nhập, đăng xuất, route protection, current-user context trong shell.",
        "Thiết lập nền tảng dữ liệu cho Story 1.2 với Prisma/PostgreSQL cho `User`, `Session`, `AuditLog`.",
        "Bổ sung audit log cho hành động đăng nhập và đăng xuất; áp dụng nguyên tắc fail-closed cho protected routes.",
        "Thiết lập kiểm tra nền tảng bằng smoke test và các bước verify build, lint, test cho phần foundation/auth."
    ]
    for item in completed:
        add_bullet(document, item)

    add_heading(document, "5. Những bước sẽ triển khai trong thời gian tới")
    upcoming = [
        "Hoàn thiện Story 1.3: quản lý người dùng, vai trò và phạm vi đơn vị để đáp ứng kiểm soát truy cập theo tổ chức.",
        "Hoàn thiện Story 1.4: permission primitives, danh mục dùng chung và cấu hình nền làm cơ sở cho các workflow nghiệp vụ.",
        "Triển khai các workflow đề tài thực sự ở backend: tiếp nhận hồ sơ, bổ sung hồ sơ, phân công đánh giá, phê duyệt và các trạng thái chuyển tiếp có kiểm soát.",
        "Kết nối storage và hạ tầng nền gồm MinIO cho tệp đính kèm, Redis cho nhắc việc và notification jobs.",
        "Phát triển các module tiếp theo của phase 1: theo dõi đề tài đã duyệt, báo cáo tiến độ, dashboard mở rộng, báo cáo và xuất Excel/PDF.",
        "Hoàn thiện môi trường triển khai nội bộ với Docker Compose, Nginx, migration và seed trên môi trường dev/staging."
    ]
    for item in upcoming:
        add_bullet(document, item)

    add_heading(document, "6. Hình ảnh giao diện hiện tại của project")
    add_body(
        document,
        "Các hình dưới đây được chụp từ giao diện hiện tại của web app trong môi trường local, thể hiện bộ khung và các màn hình nghiệp vụ đã có."
    )

    screenshots = [
        (TMP_DIR / "login-page.png", "Hình 1. Màn hình đăng nhập nội bộ"),
        (TMP_DIR / "dashboard-page.png", "Hình 2. Dashboard lãnh đạo"),
        (TMP_DIR / "proposals-page.png", "Hình 3. Danh sách hồ sơ đề tài"),
        (TMP_DIR / "proposal-detail-page.png", "Hình 4. Chi tiết hồ sơ đề tài"),
        (TMP_DIR / "tasks-page.png", "Hình 5. Danh sách giao việc"),
    ]

    for index, (image_path, caption) in enumerate(screenshots):
        if index and index % 2 == 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        add_image(document, image_path, 16.5, caption)

    add_heading(document, "7. Kiến nghị")
    add_body(
        document,
        "Đề nghị tiếp tục đầu tư triển khai đề tài theo lộ trình phase 1 đã xác lập. Với nền tảng hiện tại, project đã vượt qua giai đoạn ý tưởng và đã có cơ sở kỹ thuật đủ rõ để phát triển tiếp thành hệ thống vận hành nội bộ chính thức."
    )
    add_body(
        document,
        "Trong ngắn hạn, nên ưu tiên hoàn tất lớp phân quyền theo đơn vị và các workflow đề tài cốt lõi để sớm đưa hệ thống vào dùng thử nghiệp vụ tại phạm vi hẹp."
    )

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
